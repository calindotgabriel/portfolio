#!/usr/bin/env python3
"""Generate the Cognizant Technical Profile .docx for Calin Gabriel.

Output: cognizant/Technical Profile - Calin Gabriel.docx

Layout is inspired by the Cognizant "Example 21" template but built as a clean,
printable Word document (not a pixel-perfect clone). Sections and ordering
match the template so it can be pasted into Cognizant's real template file if
they insist on their exact styling.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = REPO_ROOT / "cognizant" / "Technical Profile - Calin Gabriel.docx"

ACCENT = RGBColor(0x00, 0x3B, 0x6F)
MUTED = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x11, 0x11, 0x11)
LIGHT_BORDER = "BFBFBF"


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        spec = kwargs[edge]
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        for key, value in spec.items():
            el.set(qn(f"w:{key}"), value)


def _add_run(paragraph, text, *, bold=False, italic=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, *, size=14, color=ACCENT, space_before=12, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    _add_run(p, text.upper(), bold=True, size=size, color=color)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "003B6F")
    borders.append(bottom)
    p_pr.append(borders)
    return p


def add_paragraph(doc, text, *, size=10, color=BLACK, space_after=4, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_bullets(doc, items, *, size=10, color=BLACK):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        # If item is a (label, body) tuple, bold the label.
        if isinstance(item, tuple):
            label, body = item
            _add_run(p, f"{label}: ", bold=True, size=size, color=color)
            _add_run(p, body, size=size, color=color)
        else:
            _add_run(p, item, size=size, color=color)


def dot_string(filled, total=5):
    return ("●" * filled) + ("○" * (total - filled))


def add_skills_table(doc, groups):
    """groups: list of (dots_int, [items])"""
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.6)
    table.columns[1].width = Cm(13.4)
    for dots, header, items in groups:
        row = table.add_row().cells
        row[0].width = Cm(3.6)
        row[1].width = Cm(13.4)
        # Left cell: dots + label
        left_p = row[0].paragraphs[0]
        _add_run(left_p, dot_string(dots) + "\n", bold=True, size=12, color=ACCENT)
        _add_run(left_p, header, bold=True, size=9, color=MUTED)
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Right cell: items
        right_p = row[1].paragraphs[0]
        _add_run(right_p, "  ·  ".join(items), size=10, color=BLACK)
        right_p.paragraph_format.space_after = Pt(0)
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Bottom border for row separation
        for cell in row:
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": LIGHT_BORDER, "space": "0"})


def add_experience_entry(doc, company, role, period, location, description, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    _add_run(p, f"{company} — ", bold=True, size=11, color=BLACK)
    _add_run(p, role, bold=True, size=11, color=ACCENT)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    _add_run(meta, f"{period}  ·  {location}", italic=True, size=9, color=MUTED)
    body = doc.add_paragraph()
    body.paragraph_format.space_after = Pt(2)
    _add_run(body, description, size=10, color=BLACK)
    tech_p = doc.add_paragraph()
    tech_p.paragraph_format.space_after = Pt(6)
    _add_run(tech_p, "Tech: ", bold=True, size=9, color=MUTED)
    _add_run(tech_p, tech, size=9, color=BLACK)


def add_project_card(doc, *, title, header_line, description, role, responsibilities, tech_stack):
    add_heading(doc, title, size=12, color=ACCENT, space_before=10, space_after=2)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    _add_run(meta, header_line, italic=True, size=9, color=MUTED)

    desc = doc.add_paragraph()
    desc.paragraph_format.space_after = Pt(4)
    _add_run(desc, description, size=10, color=BLACK)

    role_p = doc.add_paragraph()
    role_p.paragraph_format.space_after = Pt(2)
    _add_run(role_p, "Role: ", bold=True, size=10, color=BLACK)
    _add_run(role_p, role, size=10, color=BLACK)

    resp_head = doc.add_paragraph()
    resp_head.paragraph_format.space_after = Pt(2)
    _add_run(resp_head, "Key responsibilities:", bold=True, size=10, color=BLACK)
    add_bullets(doc, responsibilities)

    stack_head = doc.add_paragraph()
    stack_head.paragraph_format.space_before = Pt(4)
    stack_head.paragraph_format.space_after = Pt(2)
    _add_run(stack_head, "Tech stack:", bold=True, size=10, color=BLACK)
    add_bullets(doc, tech_stack)


# ---------- Content ----------

NAME = "Calin Gabriel"
TITLE = "Senior Backend Engineer  ·  Full-Stack (Node.js / TypeScript)"
CONTACT = [
    ("Location", "Cluj-Napoca, Romania  ·  Remote EU"),
    ("Email", "contact@calingabriel.com"),
    ("Phone", "+40 759 407 066"),
    ("LinkedIn", "linkedin.com/in/calingabriel-ts-dev"),
    ("Website", "calingabriel.com"),
]

SUMMARY = (
    "I am a senior backend engineer with around 7 years of commercial Node.js and 5+ "
    "years of TypeScript, working with Western clients in regulated fintech, energy "
    "data, and large-traffic marketplaces. I design and build Node.js / TypeScript "
    "services on NestJS, Fastify, and GraphQL, and I'm comfortable in the React "
    "front-ends that consume them. My work has centered on high-volume operational "
    "flows, Java-to-NestJS migrations backed by integration and E2E tests, and "
    "hands-on API performance work — including MongoDB query profiling that cut key "
    "response times by 65%, and a Node.js Worker-Pool refactor that reduced a "
    "45-minute export to 12 minutes. I ship through CI/CD on Azure DevOps and GitHub "
    "Actions, and I'm familiar with the containerized delivery, Kubernetes, and "
    "monitoring around it. I collaborate well in cross-functional Agile Scrum teams "
    "and own features end-to-end. I use GitHub Copilot and Claude Code day-to-day "
    "for boilerplate and debugging support."
)

AREAS = [
    ("Web Development", "Node.js, TypeScript, JavaScript"),
    ("Backend Frameworks", "NestJS, Fastify, Express.js"),
    ("Frontend", "React, Redux, Next.js"),
    ("APIs", "REST (OpenAPI / Swagger), GraphQL"),
    ("DevOps", "Azure DevOps pipelines, GitHub Actions, Docker, AWS (Lambda, S3, SQS, serverless); familiar with Kubernetes / AKS via existing Helm and Terraform config"),
    ("Databases", "MongoDB, PostgreSQL, MySQL, Redis; query profiling"),
    ("Real-time & Integrations", "WebSockets, third-party API integrations, marketing automation (Iterable)"),
    ("Testing", "Jest, Vitest, Testcontainers, LocalStack; integration + E2E + TDD"),
    ("Monitoring", "Grafana, Prometheus, DataDog"),
    ("AI Tooling", "GitHub Copilot, Claude Code"),
    ("Project Environments", "Enterprise, regulated industries, late-stage startups"),
    ("Soft skills", "Cross-functional collaboration, mentoring, code review, architectural thinking"),
]

INDUSTRIES = [
    "Regulated Fintech / Digital-Asset Custody (bank-facing)",
    "Energy & Utilities (operational data platforms)",
    "Real Estate / Large-Traffic Marketplaces",
    "Payments / E-commerce",
    "Enterprise SaaS",
]

SKILLS_GROUPS = [
    (5, "Deep senior expertise", [
        "Node.js", "TypeScript", "Backend architecture & microservices", "REST APIs (OpenAPI)",
    ]),
    (4, "Strong senior-level mastery", [
        "NestJS", "Fastify", "GraphQL", "MongoDB", "React",
        "AWS (Lambda, S3, SQS, serverless)",
        "Integration testing / TDD", "Query profiling & performance",
        "Code review & mentoring", "Cross-functional Agile Scrum",
        "GitHub Copilot", "Claude Code",
    ]),
    (3, "Solid mid-level", [
        "PostgreSQL / MySQL", "Redis",
        "Azure DevOps (CI/CD)", "GitHub Actions", "Docker",
        "Redux / Next.js", "Java (background + migrations)",
        "WebSockets / real-time",
    ]),
    (2, "Junior / practical exposure", [
        "Kubernetes / AKS",
        "Helm (consumer of existing charts)",
        "Terraform (consumer of existing config)",
        "Grafana / Prometheus",
        "Elasticsearch",
        "Playwright E2E",
    ]),
]

EXPERIENCE = [
    dict(
        company="Bitpanda",
        role="Senior Backend Developer",
        period="Sep 2025 – Present",
        location="Remote (Austria-based client)",
        description=(
            "Backend feature ownership on a bank-facing digital-asset custody platform. "
            "Cross-functional Agile Scrum team; introduced backend testing patterns "
            "adopted by the team."
        ),
        tech="Fastify, GraphQL, TypeScript, Node.js, PostgreSQL, AWS (Lambda, serverless), GitHub Actions.",
    ),
    dict(
        company="ImmoScout24",
        role="Senior Full Stack Developer",
        period="Aug 2024 – Jun 2025",
        location="Remote (Austria / Germany)",
        description=(
            "Search, listing, SEO, and lifecycle-messaging work in a mature AWS + MongoDB "
            "marketplace platform serving DACH markets under strict SEO, test-coverage, "
            "and release-safety gates."
        ),
        tech="React, Redux, Node.js, Express.js, TypeScript, MongoDB, Redis, AWS, GitHub Actions, Iterable.",
    ),
    dict(
        company="RWE",
        role="Full Stack Developer",
        period="Sep 2022 – Jul 2024",
        location="Remote (Germany)",
        description=(
            "Java → NestJS modernization for an energy-market operational data platform. "
            "Shipped service changes through Azure DevOps pipelines and worked inside an "
            "AKS + Helm + Terraform delivery setup."
        ),
        tech="NestJS, Node.js, TypeScript, React, Java (legacy), MongoDB, Redis, Azure DevOps, AKS, Helm, Terraform, Grafana, Prometheus.",
    ),
    dict(
        company="Endava",
        role="Full Stack Developer",
        period="Sep 2019 – Jul 2021",
        location="Cluj-Napoca, Romania",
        description=(
            "Node.js document-processing pipeline for a scientific-publishing client and "
            "React + TypeScript checkout with Klarna and card-payment integration for a "
            "retail client. Mentored two junior developers."
        ),
        tech="Node.js, TypeScript, React, Redux, Express.js, MongoDB, PostgreSQL, AWS (S3, Lambda, SQS), GitHub Actions.",
    ),
    dict(
        company="WIP Romania",
        role="Full Stack Developer",
        period="Jul 2018 – Sep 2019",
        location="Cluj-Napoca, Romania",
        description=(
            "Real-time mobile wallet with deposits, withdrawals, balance updates, and "
            "transaction history over WebSockets."
        ),
        tech="MeteorJS, React, WebSockets, MongoDB.",
    ),
    dict(
        company="DeverSoft",
        role="Web Developer",
        period="Oct 2013 – Mar 2015",
        location="Cluj-Napoca, Romania",
        description=(
            "Restaurant-management frontend with JavaScript + Java integration and "
            "real-time WebSocket updates across order, table, kitchen, and service workflows."
        ),
        tech="JavaScript, Java, WebSockets.",
    ),
]

EDUCATION = "BS in Computer Science  ·  Babeș-Bolyai University, Cluj-Napoca, Romania"

LANGUAGES = [
    ("Romanian", "Native"),
    ("English", "C1 fluent"),
]

OUTSIDE_OF_WORK = (
    "I treat engineering as a continuous-learning discipline — I follow a curated set "
    "of engineering voices, sharpen ideas with peers and with AI assistants, and build "
    "small proof-of-concept projects on weekends to try things out before bringing "
    "them into production. Outside of engineering I focus on family, the outdoors, "
    "and physical training — the anchors that keep energy steady across long delivery "
    "cycles."
)

PROJECTS = [
    dict(
        title="1. Institutional Digital-Asset Custody Platform",
        header_line="Regulated Fintech  ·  Remote (Austria)  ·  Enterprise",
        description=(
            "Custody platform used by banks to manage the lifecycle of institutional "
            "digital-asset holdings. High-volume, audit-critical operational flows under "
            "strict regulatory, uptime, and security requirements."
        ),
        role="Senior Backend Engineer",
        responsibilities=[
            "Owned the crypto address-book backend end-to-end: API contract design, data-model design, request/response validation, cross-service error handling, retry strategies, and structured logging.",
            "Designed and built microservices with clean domain boundaries and observability in place (logs, metrics, structured events) so each service was independently deployable and safe to evolve.",
            "Strengthened authentication and authorization across the service mesh; introduced integration-testing patterns (Testcontainers-style) that the team adopted for new services.",
            "Contributed to the GitHub Actions CI/CD flow: build, unit + integration tests, quality gates, and staged deployments to AWS.",
            "Collaborated cross-functionally with backend, frontend, product, and design in Agile Scrum; participated in code reviews and onboarded a new joiner into the crypto-custody domain.",
            "Used GitHub Copilot and Claude Code day-to-day for boilerplate scaffolding (DTOs, validators, controllers) and to accelerate debugging of async and third-party integration issues.",
        ],
        tech_stack=[
            ("Languages", "TypeScript, JavaScript"),
            ("Frameworks", "Fastify, GraphQL, Node.js"),
            ("Data Storage", "PostgreSQL"),
            ("Documentation", "OpenAPI (OAS3), GraphQL Schema"),
            ("Tools", "git, VSCode, Docker"),
            ("Cloud / DevOps", "AWS (Lambda, serverless), GitHub Actions"),
            ("AI Tooling", "GitHub Copilot, Claude Code"),
            ("Project Management", "JIRA"),
        ],
    ),
    dict(
        title="2. Energy-Market Operational Data Platform",
        header_line="Energy & Utilities  ·  Remote (Germany)  ·  Enterprise",
        description=(
            "Operational platform that ingests, processes, and reports data across a large "
            "European energy portfolio — supporting pricing, forecasting, and audit-ready "
            "reporting for energy trading and grid operations. Data pipelines with strict "
            "SLAs, uptime, and traceability constraints."
        ),
        role="Full Stack Engineer (backend-heavy)",
        responsibilities=[
            "Led incremental Java → NestJS microservice migration using a strangler-fig pattern; wrote integration and E2E tests that made regression risk visible before each cutover and protected production behavior during the transition.",
            "Profiled MongoDB aggregations serving 15,000+ daily queries; introduced indexing, projection tightening, and query refactors that reduced key API response times by 65%.",
            "Moved CPU-intensive export workloads out of the request path into a Node.js Worker-Pool architecture, cutting export runtimes from 45 minutes to 12 minutes while keeping normal API traffic responsive under peak load.",
            "Shipped service changes through Azure DevOps CI/CD pipelines (multi-stage build, test, staged deployment) into an AKS cluster; worked with the existing Helm charts and Terraform configuration when a change touched the delivery layer — comfortable consuming this setup, not the primary IaC author.",
            "Used Grafana and Prometheus dashboards during incidents to trace issues back to service changes and to inform controlled hotfixes.",
            "Used GitHub Copilot for boilerplate and refactoring; wrote and updated service documentation via NSwag / OpenAPI for cross-team integration.",
            "Collaborated with QA, architects, product owners, and business stakeholders to refine requirements, clarify edge cases, and align technical implementation with business goals; contributed to code review and mentored newer engineers on testing and migration practices.",
        ],
        tech_stack=[
            ("Languages", "TypeScript, JavaScript, Java (legacy paths)"),
            ("Frameworks", "NestJS, Node.js, React"),
            ("Data Storage", "MongoDB, Redis"),
            ("Documentation", "OpenAPI (OAS3), NSwag"),
            ("Async patterns", "HTTP-driven services + scheduled jobs (cron)"),
            ("Tools", "git, JetBrains IDEs, Docker"),
            ("Cloud / DevOps", "Azure DevOps pipelines, AKS (consumer), Helm (consumer of existing charts), Terraform (consumer of existing config)"),
            ("Monitoring", "Grafana, Prometheus"),
            ("AI Tooling", "GitHub Copilot"),
            ("Project Management", "JIRA"),
        ],
    ),
    dict(
        title="3. Large-Traffic Marketplace Platform",
        header_line="PropTech  ·  Remote (Austria / Germany)  ·  Enterprise",
        description=(
            "Real-estate marketplace serving millions of monthly users across DACH markets. "
            "Supports search, listing lifecycle, and lifecycle-messaging integration under "
            "strict SEO, test-coverage, and production-safety standards."
        ),
        role="Senior Full Stack Engineer",
        responsibilities=[
            "Integrated Iterable marketing automation into React and Node.js product flows for targeted lifecycle messaging, satisfying SEO, test-coverage, and release-safety gates.",
            "Shipped search, listing, and SEO improvements across React and Node.js in a mature AWS + MongoDB environment.",
            "Matched the platform's established review, integration-test, and production-safety standards while working inside a distributed system.",
            "Collaborated with product, design, and analytics counterparts to turn lifecycle-messaging requirements into concrete backend contracts.",
            "Used GitHub Copilot and Claude Code for refactoring and for test scaffolding when moving inside a large monorepo.",
        ],
        tech_stack=[
            ("Languages", "TypeScript, JavaScript"),
            ("Frameworks", "React, Redux, Node.js, Express.js"),
            ("Data Storage", "MongoDB, Redis"),
            ("Documentation", "OpenAPI (OAS3)"),
            ("Third-party", "Iterable (marketing automation)"),
            ("Tools", "git, VSCode, Docker"),
            ("Cloud / DevOps", "AWS, GitHub Actions"),
            ("AI Tooling", "GitHub Copilot, Claude Code"),
            ("Project Management", "JIRA"),
        ],
    ),
    dict(
        title="4. Enterprise Document Processing & Payments",
        header_line="Publishing (documents) + E-commerce (payments)  ·  Cluj-Napoca, Romania  ·  Enterprise",
        description=(
            "Two enterprise workstreams delivered in tandem: (a) a Node.js pipeline for a "
            "scientific-publishing client processing 2,000+ documents per day across AWS "
            "S3 and Lambda; (b) a React + TypeScript checkout with Klarna and "
            "card-payment integration for a retail client. Both required audit-ready "
            "processing and third-party integration under enterprise release standards."
        ),
        role="Full Stack Engineer",
        responsibilities=[
            "Built a Node.js pipeline that processed 2,000+ scientific papers daily on AWS S3 and Lambda, with structured logging and retry semantics that made failure modes visible and recoverable.",
            "Delivered React submission and tracking flows for large document collections and collaborative editorial workflows.",
            "Shipped TypeScript checkout support integrating Klarna and card-payment providers; validated end-to-end via integration tests against real payment sandboxes.",
            "Mentored two junior developers on Node.js, TypeScript, and testing patterns; contributed to code-review culture inside the team.",
            "Contributed to the GitHub Actions CI/CD flow for build, test, and controlled deployment to AWS.",
        ],
        tech_stack=[
            ("Languages", "TypeScript, JavaScript"),
            ("Frameworks", "React, Redux, Node.js, Express.js"),
            ("Data Storage", "MongoDB, PostgreSQL"),
            ("Documentation", "OpenAPI (OAS3)"),
            ("Messaging", "AWS SQS"),
            ("Third-party", "Klarna, card-payment providers"),
            ("Tools", "git, VSCode, Docker"),
            ("Cloud / DevOps", "AWS (S3, Lambda), GitHub Actions"),
            ("Project Management", "JIRA"),
        ],
    ),
]


def build():
    doc = Document()

    # Base font
    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Header (name + title + contact)
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(0)
    _add_run(name_p, NAME, bold=True, size=22, color=BLACK)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    _add_run(title_p, TITLE, italic=True, size=11, color=ACCENT)

    contact_tbl = doc.add_table(rows=len(CONTACT), cols=2)
    contact_tbl.autofit = False
    contact_tbl.columns[0].width = Cm(3.2)
    contact_tbl.columns[1].width = Cm(13.8)
    for idx, (label, value) in enumerate(CONTACT):
        left = contact_tbl.rows[idx].cells[0]
        right = contact_tbl.rows[idx].cells[1]
        left.width = Cm(3.2)
        right.width = Cm(13.8)
        _add_run(left.paragraphs[0], label, bold=True, size=9, color=MUTED)
        _add_run(right.paragraphs[0], value, size=10, color=BLACK)

    # Summary
    add_heading(doc, "Summary")
    add_paragraph(doc, SUMMARY, space_after=6)

    # Areas of expertise
    add_heading(doc, "Areas of expertise")
    add_bullets(doc, AREAS)

    # Industries
    add_heading(doc, "Industries")
    add_bullets(doc, INDUSTRIES)

    # Skills and proficiency
    add_heading(doc, "Skills and proficiency")
    add_paragraph(
        doc,
        "● filled = mastered  ·  ○ empty = to grow  ·  5-dot deep expertise, 4-dot strong senior, 3-dot solid IC, 2-dot practical exposure",
        size=9, color=MUTED, space_after=6,
    )
    add_skills_table(doc, SKILLS_GROUPS)

    # Work experience
    add_heading(doc, "Work experience")
    for entry in EXPERIENCE:
        add_experience_entry(doc, **entry)

    # Education
    add_heading(doc, "Education")
    add_paragraph(doc, EDUCATION, space_after=4)

    # Languages
    add_heading(doc, "Languages")
    for lang, level in LANGUAGES:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _add_run(p, f"{lang} — ", bold=True, size=10, color=BLACK)
        _add_run(p, level, size=10, color=BLACK)

    # Outside of work
    add_heading(doc, "Outside of work")
    add_paragraph(doc, OUTSIDE_OF_WORK, space_after=6)

    # Selected relevant project experience
    add_heading(doc, "Selected relevant project experience")
    for project in PROJECTS:
        add_project_card(doc, **project)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
