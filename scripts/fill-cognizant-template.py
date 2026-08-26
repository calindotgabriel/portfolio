#!/usr/bin/env python3
"""Fill the Cognizant Technical Profile template with Calin's content.

Reads `cognizant/Technical Profile - Example 21.docx` (the reference template),
substitutes Calin's content into the existing paragraphs and cells, and writes
`cognizant/Technical Profile - Calin Gabriel [filled].docx`.

Formatting is preserved because we only modify text runs — never paragraph or
cell properties. When a cell needs more lines than the template has, we clone
the last paragraph's XML so the new line inherits its style.

Contact-info cells (Email/Phone/LinkedIn) have their label on the same
paragraph as the eventual value. We append a fresh paragraph after the label
with the actual value.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document

REPO = Path(__file__).resolve().parents[1]
SRC = REPO / "cognizant" / "Technical Profile - Example 21.docx"
DST = REPO / "cognizant" / "Technical Profile - Calin Gabriel [filled].docx"


# ---------- helpers ----------

def set_paragraph_text(paragraph, text):
    """Set paragraph text, preserving the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def clone_paragraph_after(paragraph, text):
    """Clone `paragraph`'s XML, insert it after, and set the new paragraph's text.

    Returns the new paragraph object.
    """
    new_p_xml = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p_xml)
    # Walk the parent to find the wrapped Paragraph object again
    parent = paragraph._parent
    for p in parent.paragraphs:
        if p._p is new_p_xml:
            set_paragraph_text(p, text)
            return p
    return None


def replace_cell_lines(cell, lines, *, keep_first=False):
    """Replace all paragraphs in `cell` with `lines`.

    - keep_first=True preserves the first paragraph (typically a header like
      "Areas of expertise") and replaces the rest.
    - Extra existing paragraphs (beyond len(lines)) are emptied.
    - If we need more lines than exist, we clone the last paragraph.
    """
    existing = list(cell.paragraphs)
    start = 1 if keep_first else 0
    slots = existing[start:]

    # Reuse existing slots
    for idx, text in enumerate(lines):
        if idx < len(slots):
            set_paragraph_text(slots[idx], text)
        else:
            # Need to clone the last paragraph in the cell
            template = cell.paragraphs[-1]
            clone_paragraph_after(template, text)

    # Empty excess existing paragraphs
    if len(lines) < len(slots):
        for extra in slots[len(lines):]:
            set_paragraph_text(extra, "")


def find_paragraph_by_prefix(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


# ---------- Calin content ----------

CONTACT = {
    "Email": "contact@calingabriel.com",
    "Phone": "+40 759 407 066",
    "Linkedin link": "linkedin.com/in/calingabriel-ts-dev",
}

SUMMARY = (
    "I am a senior backend engineer with around 7 years of commercial Node.js "
    "and 5+ years of TypeScript, working with Western clients in regulated "
    "fintech, energy data, and large-traffic marketplaces. I design and build "
    "Node.js / TypeScript services on NestJS, Fastify, and GraphQL, and I'm "
    "comfortable in the React front-ends that consume them. My work has "
    "centered on high-volume operational flows, Java-to-NestJS migrations "
    "backed by integration and E2E tests, and hands-on API performance work "
    "— including MongoDB query profiling that cut key response times by 65%, "
    "and a Node.js Worker-Pool refactor that reduced a 45-minute export to "
    "12 minutes. I ship through CI/CD on Azure DevOps and GitHub Actions, "
    "and I'm familiar with the containerized delivery, Kubernetes, and "
    "monitoring around it. I use GitHub Copilot and Claude Code day-to-day "
    "for boilerplate and debugging support."
)

SUMMARY_LINE_2 = (
    "I collaborate well in cross-functional Agile Scrum teams and own "
    "features end-to-end. I've delivered for regulated fintech (Bitpanda), "
    "energy operations (RWE), large-traffic marketplaces (ImmoScout24), and "
    "enterprise document/payment platforms (Endava)."
)

AREAS_OF_EXPERTISE = [
    "Web Development: Node.js, TypeScript, JavaScript",
    "Backend Frameworks: NestJS, Fastify, Express.js",
    "Frontend: React, Redux, Next.js",
    "APIs: REST (OpenAPI / Swagger), GraphQL",
    "DevOps: Azure DevOps pipelines, GitHub Actions, Docker, AWS (Lambda, S3, SQS, serverless); Kubernetes / AKS, Helm, Terraform at working level",
    "Databases: MongoDB, PostgreSQL, MySQL, Redis; query profiling",
    "Real-time & Integrations: WebSockets, third-party API integrations, marketing automation (Iterable)",
    "Testing: Jest, Vitest, Testcontainers, LocalStack; integration + E2E + TDD",
    "Monitoring: Grafana, Prometheus, DataDog",
    "AI Tooling: GitHub Copilot, Claude Code",
    "Project Environments: Enterprise, regulated industries, late-stage startups",
    "Soft skills: Cross-functional collaboration, mentoring, code review, architectural thinking",
]

INDUSTRIES = [
    "Regulated Fintech / Digital-Asset Custody",
    "Energy & Utilities",
    "Real Estate / Marketplaces",
    "Payments / E-commerce",
    "Enterprise SaaS",
]

# Work experience: 4 job blocks, following the template's structure
# Each block has: role title + company + responsibilities + tech
EXPERIENCE = [
    dict(
        role="Senior Backend Developer",
        company="Bitpanda",
        period="Sep 2025 – Present",
        bullets=[
            "Backend feature ownership on a bank-facing digital-asset custody platform used by institutional clients",
            "Owned the crypto address-book service end-to-end: API contract design, data-model, validation, tests, and AWS deployment",
            "Introduced backend integration-testing patterns (Testcontainers-style) later adopted by the team",
            "Strengthened authentication and authorization across services",
        ],
        tech="Tech Stack: Fastify, GraphQL, TypeScript, Node.js, PostgreSQL, AWS (Lambda, serverless), GitHub Actions, GitHub Copilot, Claude Code",
    ),
    dict(
        role="Senior Full Stack Developer",
        company="ImmoScout24",
        period="Aug 2024 – Jun 2025",
        bullets=[
            "Integrated Iterable marketing automation into React and Node.js flows for targeted lifecycle messaging",
            "Shipped search, listing, and SEO improvements in a mature AWS + MongoDB marketplace platform",
            "Matched the platform's established review, integration-test, and production-safety standards",
        ],
        tech="Tech Stack: React, Redux, Node.js, Express.js, TypeScript, MongoDB, Redis, AWS, GitHub Actions, Iterable, GitHub Copilot, Claude Code",
    ),
    dict(
        role="Full Stack Developer",
        company="RWE",
        period="Sep 2022 – Jul 2024",
        bullets=[
            "Led incremental Java → NestJS microservice migration for an energy-market operational data platform",
            "Profiled MongoDB aggregations serving 15,000+ daily queries; cut key API response times by 65%",
            "Refactored CPU-intensive exports into a Node.js Worker-Pool architecture, reducing runtime from 45 to 12 minutes",
            "Shipped through Azure DevOps CI/CD pipelines into an AKS cluster; worked with the existing Helm and Terraform config",
        ],
        tech="Tech Stack: NestJS, Node.js, TypeScript, React, Java (legacy), MongoDB, Redis, Azure DevOps, AKS, Helm, Terraform, Grafana, Prometheus, GitHub Copilot",
    ),
    dict(
        role="Full Stack Developer",
        company="Endava",
        period="Sep 2019 – Jul 2021",
        bullets=[
            "Built a Node.js pipeline processing 2,000+ scientific papers daily on AWS S3 and Lambda",
            "Delivered React submission and tracking flows for large document collections",
            "Shipped TypeScript checkout support integrating Klarna and card-payment providers",
            "Mentored two junior developers on Node.js, TypeScript, and testing patterns",
        ],
        tech="Tech Stack: Node.js, TypeScript, React, Redux, Express.js, MongoDB, PostgreSQL, AWS (S3, Lambda, SQS), GitHub Actions",
    ),
]

EDUCATION = dict(
    left=["Computer Science, Babeș-Bolyai University, Romania", "2014 – 2017"],
    right=["Bachelor's degree in computer science"],
)

LANGUAGES = ["English", "Romanian"]  # already in template

SKILLS_PANEL = [
    "Backend architecture",
    "APIs (OpenAPI, GraphQL)",
    "Node.js | TypeScript",
    "NestJS | Fastify",
    "React | Redux",
    "MongoDB | PostgreSQL",
    "AWS | Azure DevOps",
    "Cross-functional teamwork",
    "Agile | Scrum",
    "Code review | Mentoring",
]

OUTSIDE_OF_WORK_LINE_1 = (
    "I treat engineering as a continuous-learning discipline — I follow a "
    "curated set of engineering voices, sharpen ideas with peers and with AI "
    "assistants, and build small proof-of-concept projects on weekends to try "
    "things out before bringing them into production."
)
OUTSIDE_OF_WORK_LINE_2 = (
    "Outside of engineering I focus on family, the outdoors, and physical "
    "training — the anchors that keep energy steady across long delivery cycles."
)

# Project cards for Table 5 (5 project slots in the template — we fill 4 and blank the 5th)
PROJECTS = [
    dict(
        left_top="Bitpanda / Regulated Fintech",
        left_bottom="Remote (Austria)",
        left_scale="Enterprise",
        title="Institutional Digital-Asset Custody Platform",
        role="Role: Senior Backend Engineer",
        bullets=[
            "Own the crypto address-book backend end-to-end: API contract, data-model, validation, error handling, retries, structured logging",
            "Design microservices with clean domain boundaries and observability (logs, metrics, structured events)",
            "Strengthen authentication and authorization across the service mesh",
            "Introduce integration-testing patterns (Testcontainers-style) that the team adopted for new services",
            "Ship through GitHub Actions CI/CD: build, unit + integration, quality gates, staged deploys to AWS",
            "Use GitHub Copilot and Claude Code day-to-day for boilerplate and debugging",
            "Cross-functional Agile Scrum team; participate in code reviews and onboard new joiners",
        ],
        tech_lines=[
            "Languages: TypeScript, JavaScript",
            "Frameworks: Fastify, GraphQL, Node.js",
            "Data Storage: PostgreSQL",
            "Documentation: OpenAPI (OAS3), GraphQL Schema",
            "Tools: git, VSCode, Docker",
            "Cloud/DevOps: AWS (Lambda, serverless), GitHub Actions",
            "AI Tooling: GitHub Copilot, Claude Code",
            "Project Management: JIRA",
        ],
    ),
    dict(
        left_top="RWE / Energy & Utilities",
        left_bottom="Remote (Germany)",
        left_scale="Enterprise",
        title="Energy-Market Operational Data Platform",
        role="Role: Full Stack Engineer (backend-heavy)",
        bullets=[
            "Lead incremental Java → NestJS microservice migration using a strangler-fig pattern, with integration and E2E test harnesses",
            "Profile MongoDB aggregations serving 15,000+ daily queries; reduce key API response times by 65%",
            "Move CPU-intensive export workloads into a Node.js Worker-Pool architecture: 45 min → 12 min while keeping API responsive",
            "Ship service changes through Azure DevOps CI/CD pipelines into AKS; work with existing Helm and Terraform config",
            "Use Grafana and Prometheus during incidents to trace issues back to service changes",
            "Use GitHub Copilot for boilerplate and refactoring; document services via NSwag / OpenAPI",
            "Collaborate with QA, architects, POs, and stakeholders; contribute to code review and mentor newer engineers",
        ],
        tech_lines=[
            "Languages: TypeScript, JavaScript, Java (legacy)",
            "Frameworks/Libraries: NestJS, Node.js, React",
            "Data Storage: MongoDB, Redis",
            "Documentation: OpenAPI (OAS3), NSwag",
            "Async patterns: HTTP + scheduled jobs (cron)",
            "Tools: git, JetBrains IDEs, Docker",
            "Cloud/DevOps: Azure DevOps, AKS, Helm, Terraform",
            "Monitoring: Grafana, Prometheus",
            "AI Tooling: GitHub Copilot",
            "Project Management: JIRA",
        ],
    ),
    dict(
        left_top="ImmoScout24 / PropTech",
        left_bottom="Remote (Austria / Germany)",
        left_scale="Enterprise",
        title="Large-Traffic Marketplace Platform",
        role="Role: Senior Full Stack Engineer",
        bullets=[
            "Integrate Iterable marketing automation into React and Node.js flows for targeted lifecycle messaging",
            "Ship search, listing, and SEO improvements across React and Node.js in a mature AWS + MongoDB environment",
            "Match established review, integration-test, and production-safety standards inside a distributed system",
            "Turn lifecycle-messaging requirements from product / design / analytics into concrete backend contracts",
            "Use GitHub Copilot and Claude Code for refactoring and test scaffolding in a large monorepo",
            "Follow the Agile Scrum methodology",
        ],
        tech_lines=[
            "Languages: TypeScript, JavaScript",
            "Frameworks/Libraries: React, Redux, Node.js, Express.js",
            "Data Storage: MongoDB, Redis",
            "Documentation: OpenAPI (OAS3)",
            "Third-party: Iterable (marketing automation)",
            "Tools: git, VSCode, Docker",
            "Cloud/DevOps: AWS, GitHub Actions",
            "AI Tooling: GitHub Copilot, Claude Code",
            "Project Management: JIRA",
        ],
    ),
    dict(
        left_top="Endava / Publishing + E-commerce",
        left_bottom="Cluj-Napoca, Romania",
        left_scale="Enterprise",
        title="Enterprise Document Processing & Payments",
        role="Role: Full Stack Engineer",
        bullets=[
            "Build a Node.js pipeline processing 2,000+ scientific papers daily on AWS S3 and Lambda",
            "Deliver React submission and tracking flows for large document collections and editorial workflows",
            "Ship TypeScript checkout support integrating Klarna and card-payment providers, validated against real payment sandboxes",
            "Mentor two junior developers on Node.js, TypeScript, and testing patterns",
            "Contribute to GitHub Actions CI/CD flow for build, test, and controlled deployment to AWS",
            "Follow the Agile Scrum methodology",
        ],
        tech_lines=[
            "Languages: TypeScript, JavaScript",
            "Frameworks/Libraries: React, Redux, Node.js, Express.js",
            "Data Storage: MongoDB, PostgreSQL",
            "Documentation: OpenAPI (OAS3)",
            "Messaging: AWS SQS",
            "Third-party: Klarna, card-payment providers",
            "Tools: git, VSCode, Docker",
            "Cloud/DevOps: AWS (S3, Lambda), GitHub Actions",
            "Project Management: JIRA",
        ],
    ),
]


# ---------- fill routine ----------

def fill():
    doc = Document(SRC)

    # ----- HEADER TABLE 0 -----
    header = doc.tables[0]
    # Row 0 C1: name + title
    set_paragraph_text(header.rows[0].cells[1].paragraphs[0], "Calin Gabriel")
    set_paragraph_text(header.rows[0].cells[1].paragraphs[1], "Senior Backend Engineer (Full-Stack, Node.js / TS)")
    # Row 1 / Row 2 C1 duplicate the same (usually vertically merged) — the audit
    # showed identical strings, so overwriting the first row is enough. Set the
    # others in case they're separate cells.
    for r in (1, 2):
        for pi, txt in enumerate(("Calin Gabriel", "Senior Backend Engineer (Full-Stack, Node.js / TS)")):
            if pi < len(header.rows[r].cells[1].paragraphs):
                set_paragraph_text(header.rows[r].cells[1].paragraphs[pi], txt)

    # Contact info: label paragraph exists (e.g. "  Email"). We append a new
    # paragraph after it with the value, cloning the label's paragraph for style.
    contact_map = [(0, "Email"), (1, "Phone"), (2, "Linkedin link")]
    for row_idx, label in contact_map:
        cell = header.rows[row_idx].cells[2]
        label_para = cell.paragraphs[0]
        clone_paragraph_after(label_para, CONTACT[label])

    # ----- SUMMARY (body paragraphs P4, P5) -----
    body = doc.paragraphs
    # Find the "Summary" heading and the two paragraphs after it
    for i, p in enumerate(body):
        if p.text.strip().startswith("Summary "):
            set_paragraph_text(body[i + 2], SUMMARY)  # body[i+1] is blank spacer
            if i + 3 < len(body):
                set_paragraph_text(body[i + 3], SUMMARY_LINE_2)
            break

    # ----- AREAS OF EXPERTISE + INDUSTRIES (Table 1) -----
    t1 = doc.tables[1]
    replace_cell_lines(t1.rows[0].cells[0], AREAS_OF_EXPERTISE, keep_first=True)
    replace_cell_lines(t1.rows[0].cells[1], INDUSTRIES, keep_first=True)

    # ----- WORK EXPERIENCE (Table 2) -----
    # Layout per job (3 rows): R{n} C1 = role title; R{n+1} C0 = company, C1 = bullets + tech;
    #                          R{n+2} C0 = period, C1 = bullets + tech (duplicated)
    # Jobs start at row 0, 3, 6, 9.
    t2 = doc.tables[2]
    for idx, job in enumerate(EXPERIENCE):
        role_row = idx * 3        # e.g. 0, 3, 6, 9
        company_row = role_row + 1
        period_row = role_row + 2

        # Role
        set_paragraph_text(t2.rows[role_row].cells[1].paragraphs[0], job["role"])
        # Company (left cell of next row)
        set_paragraph_text(t2.rows[company_row].cells[0].paragraphs[0], job["company"])
        # Bullets + tech (right cell of next row)
        combined = job["bullets"] + [job["tech"]]
        replace_cell_lines(t2.rows[company_row].cells[1], combined)
        # Period (left cell of period row)
        set_paragraph_text(t2.rows[period_row].cells[0].paragraphs[0], job["period"])
        # Right cell of period row is a mirror — leave it or clear it
        replace_cell_lines(t2.rows[period_row].cells[1], combined)

    # ----- EDUCATION (Table 3) -----
    t3 = doc.tables[3]
    # Row 1: uni block. C0 = school + years, C1 = degree line(s)
    replace_cell_lines(t3.rows[1].cells[0], EDUCATION["left"])
    replace_cell_lines(t3.rows[1].cells[1], EDUCATION["right"])
    # Row 3: high school block — clear it (Calin only wants BS listed)
    if len(t3.rows) > 3:
        replace_cell_lines(t3.rows[3].cells[0], [""])
        replace_cell_lines(t3.rows[3].cells[1], [""])

    # ----- LANGUAGES + SKILLS PANEL (Table 4) -----
    t4 = doc.tables[4]
    # Languages: R1 C0 = English, R2 C0 = Romanian — keep as-is
    # Skills panel: C2 has skill names in rows R1-R11. Replace with Calin's skills.
    skill_rows = list(range(1, 12))  # R1..R11 (11 slots)
    for row_i, text in zip(skill_rows, SKILLS_PANEL):
        set_paragraph_text(t4.rows[row_i].cells[2].paragraphs[0], text)
    # Clear any remaining slots we didn't fill
    for row_i in skill_rows[len(SKILLS_PANEL):]:
        set_paragraph_text(t4.rows[row_i].cells[2].paragraphs[0], "")

    # ----- OUTSIDE OF WORK (body paragraphs) -----
    for i, p in enumerate(body):
        if p.text.strip().startswith("Outside of work "):
            # Find the two content paragraphs after (skip blank spacers)
            content_indices = []
            for j in range(i + 1, min(i + 6, len(body))):
                if body[j].text.strip():
                    content_indices.append(j)
                if len(content_indices) == 2:
                    break
            if content_indices:
                set_paragraph_text(body[content_indices[0]], OUTSIDE_OF_WORK_LINE_1)
            if len(content_indices) > 1:
                set_paragraph_text(body[content_indices[1]], OUTSIDE_OF_WORK_LINE_2)
            break

    # ----- SELECTED RELEVANT PROJECT EXPERIENCE (Table 5) -----
    # Each project occupies 3 rows:
    #   header_row  : C0 = client-label + city (P0, P2); C1/C2 = project title
    #   content_row : C0 merged with header; C1 = role + bullets; C2 = tech stack
    #   scale_row   : C0 P1 = scale label (Enterprise / Startup); C1/C2 duplicate content
    # Header rows: 1, 8, 12, 16, 20 (per audit)
    t5 = doc.tables[5]
    project_start_rows = [1, 8, 12, 16, 20]
    for idx, project in enumerate(PROJECTS):
        header_row = project_start_rows[idx]
        content_row = header_row + 1
        scale_row = header_row + 2

        # Header row: client label + city (in C0), title (in C1 and C2 mirror)
        left_cell = t5.rows[header_row].cells[0]
        if len(left_cell.paragraphs) >= 1:
            set_paragraph_text(left_cell.paragraphs[0], project["left_top"])
        if len(left_cell.paragraphs) >= 3:
            set_paragraph_text(left_cell.paragraphs[2], project["left_bottom"])

        set_paragraph_text(t5.rows[header_row].cells[1].paragraphs[0], project["title"])
        set_paragraph_text(t5.rows[header_row].cells[2].paragraphs[0], project["title"])

        # Content row: middle cell = role + "Key responsibilities:" + bullets
        middle = [project["role"], "Key responsibilities:"] + project["bullets"]
        replace_cell_lines(t5.rows[content_row].cells[1], middle)

        # Content row: right cell = tech stack lines
        replace_cell_lines(t5.rows[content_row].cells[2], project["tech_lines"])

        # Scale row: C0 P1 holds the scale label
        scale_cell = t5.rows[scale_row].cells[0]
        target_pi = 1 if len(scale_cell.paragraphs) > 1 else 0
        set_paragraph_text(scale_cell.paragraphs[target_pi], project["left_scale"])
        # Scale row C1/C2 in the original mirror the content row — set them too in
        # case they're separate cells rather than vertically merged.
        replace_cell_lines(t5.rows[scale_row].cells[1], middle)
        replace_cell_lines(t5.rows[scale_row].cells[2], project["tech_lines"])

    # If the template has more slots than projects, blank the extras.
    for header_row in project_start_rows[len(PROJECTS):]:
        for row_i in (header_row, header_row + 1, header_row + 2):
            if row_i < len(t5.rows):
                for cell in t5.rows[row_i].cells:
                    replace_cell_lines(cell, [""])

    # Save
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    fill()
